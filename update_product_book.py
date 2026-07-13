from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_BREAK

SRC = "Dolcia_Product_Book_MASTER_v15_MAJ.docx"
OUT = "Dolcia_Product_Book_MASTER_v16_2030.docx"

doc = Document(SRC)

replacements = {
    "Surprends-moi avec ces critères": "Compose-moi mon programme",
    "Surprends-moi": "Compose-moi mon expérience",
    "Surprendre": "Composer",
    "surprends-moi": "compose-moi mon expérience",
    "surprendre intelligemment": "composer intelligemment",
}

def replace_runs(paragraph):
    for run in paragraph.runs:
        for old, new in replacements.items():
            if old in run.text:
                run.text = run.text.replace(old, new)

for p in doc.paragraphs:
    replace_runs(p)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_runs(p)

doc.add_page_break()
title = doc.add_heading("MISE À JOUR MASTER — DOLCIA 2030", level=1)
title.runs[0].font.color.rgb = RGBColor(176, 143, 80)
doc.add_paragraph("Version 16 • Décisions consolidées • Le Touquet, ville pilote")

doc.add_heading("1. Source unique de vérité", level=2)
doc.add_paragraph("Cette version complète le Product Book v15 sans supprimer ses fiches techniques. Elle synchronise désormais la vision, le vocabulaire, les parcours, l’identité visuelle et les règles fonctionnelles avec l’application MASTER.")

doc.add_heading("2. Direction artistique hybride et cohérente", level=2)
for text in [
    "Conserver le noir profond, le champagne, la typographie éditoriale et l’immersion émotionnelle de l’application actuelle.",
    "Récupérer du Product Book les compositions les plus désirables, la richesse des fiches et les meilleures idées photographiques.",
    "Chaque page utilise les mêmes couleurs, typographies, boutons, cartes, espacements, icônes, animations et mots d’action.",
    "Les photos doivent représenter la destination et l’activité réelles. Une image d’ambiance est autorisée uniquement si elle est clairement identifiée et cohérente avec Le Touquet ou la Côte d’Opale.",
    "Une fiche sans adresse vérifiable, sans source fiable ou avec un visuel trompeur ne doit pas être publiée.",
]: doc.add_paragraph(text, style="List Bullet")

doc.add_heading("3. Fonction signature : Compose-moi mon expérience", level=2)
doc.add_paragraph("L’ancien nom « Surprends-moi » est abandonné. La fonction signature compose un véritable programme chronologique selon la date, la durée, le moment, le budget, les participants, la météo, la distance et les goûts.")
for text in [
    "Accès générique sur l’accueil : Composer mon expérience.",
    "Libellé contextuel : Compose-moi ma journée, Compose-moi ma soirée ou Compose-moi un moment.",
    "Position visible et stratégique conformément à la maquette de référence, sans être confondue avec le renouvellement des résultats.",
    "Régénération possible du programme entier ou d’une seule activité.",
    "Programme présenté sous forme d’agenda horaire, avec ajout et retrait activité par activité.",
    "Changer cette idée renouvelle un seul créneau ; Voir d’autres idées renouvelle une sélection. Ces actions restent distinctes de la composition complète.",
]: doc.add_paragraph(text, style="List Bullet")

doc.add_heading("4. Fiches premium et photographies", level=2)
for text in [
    "Galerie jusqu’à six photos réelles lorsque la source le permet.",
    "Adresse complète, carte, distance, horaires, téléphone, site, réservation, prix, accessibilité, durée, public conseillé, météo et avis.",
    "Croisement Google Places, offices de tourisme, OpenAgenda, DATAtourisme et partenaires validés.",
    "Contrôle anti-fiche pauvre : les informations essentielles doivent être présentes avant mise en avant.",
    "Contrôle de date strict pour les événements et contrôle géographique strict autour de la destination choisie.",
]: doc.add_paragraph(text, style="List Bullet")

doc.add_heading("5. Comptes, mémoire IA et consentement", level=2)
for text in [
    "Le particulier peut découvrir sans compte, mais crée un compte pour conserver agendas, favoris, notes et préférences sur tous ses appareils.",
    "Acceptation des CGU et de la politique de confidentialité obligatoire ; consentement IA distinct, explicite et révocable.",
    "L’utilisateur peut consulter, corriger, exporter ou effacer ses données et la mémoire de personnalisation.",
    "Les cœurs, avis, notes, refus, réservations et agendas améliorent progressivement les recommandations.",
]: doc.add_paragraph(text, style="List Bullet")

doc.add_heading("6. Espace professionnel séparé", level=2)
for text in [
    "Aucun accès à l’interface partenaire depuis l’application client.",
    "Le professionnel crée un compte sur un portail dédié, dépose sa candidature et ses justificatifs, puis attend la validation Dolcia.",
    "Le cockpit, les packs, abonnements et tarifs professionnels sont invisibles aux clients.",
    "Missions, disponibilités, publications, réputation et notifications push ne deviennent actifs qu’après validation.",
    "Un partenaire payant ne remonte avant un non-payant qu’en cas d’égalité des critères utiles au consommateur.",
]: doc.add_paragraph(text, style="List Bullet")

doc.add_heading("7. Règles de cohérence transversale", level=2)
for text in [
    "La destination, la date, la durée, le budget et les participants restent actifs sur toutes les pages.",
    "Une activité ajoutée depuis une fiche apparaît immédiatement dans l’agenda.",
    "Le même libellé produit toujours la même action.",
    "Client, partenaire et administration partagent l’ADN Dolcia mais disposent d’interfaces adaptées à leurs usages et de droits totalement séparés.",
    "Le Touquet est le territoire pilote ; aucun résultat lointain ne doit être mélangé sans demande explicite d’élargissement.",
]: doc.add_paragraph(text, style="List Bullet")

doc.add_heading("8. Principe produit", level=2)
lead = doc.add_paragraph()
run = lead.add_run("Dolcia ne montre pas simplement des loisirs : Dolcia transforme un temps disponible en expérience désirable, crédible et immédiatement organisable.")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(176, 143, 80)

doc.core_properties.title = "Dolcia Product Book MASTER v16 — 2030"
doc.core_properties.subject = "Référence produit, visuelle et technique consolidée"
doc.save(OUT)
print(OUT)
