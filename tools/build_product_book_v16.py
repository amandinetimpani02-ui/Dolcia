from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Dolcia_Product_Book_MASTER_v16_CONFIDENTIEL_2026-07-15.docx"
HERO = ROOT / "assets" / "dolcia-eclat-concept.png"

BLACK = "080909"
CREAM = "F4EFE4"
GOLD = "D8BC7C"
MUTED = "8E877A"
WHITE = "FFFFFF"
INK = "171512"


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(section):
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def field(paragraph, instruction):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = instruction
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.extend([fldChar1, instrText, fldChar2])


doc = Document()
section = doc.sections[0]
margins(section)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)
styles["Normal"].font.color.rgb = RGBColor.from_string(INK)
styles["Normal"].paragraph_format.space_after = Pt(7)
styles["Title"].font.name = "Georgia"
styles["Title"].font.size = Pt(42)
styles["Title"].font.color.rgb = RGBColor.from_string(CREAM)
for name, size in (("Heading 1", 28), ("Heading 2", 17), ("Heading 3", 12)):
    styles[name].font.name = "Georgia"
    styles[name].font.size = Pt(size)
    styles[name].font.color.rgb = RGBColor.from_string(INK)
    styles[name].font.bold = False
    styles[name].paragraph_format.space_before = Pt(6)
    styles[name].paragraph_format.space_after = Pt(8)

for sec in doc.sections:
    margins(sec)
    hp = sec.header.paragraphs[0]
    hp.text = "DOLCIA  •  CONFIDENTIEL — DIFFUSION CONTRÔLÉE  •  MASTER v16"
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in hp.runs:
        r.font.name = "Aptos"
        r.font.size = Pt(7.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(GOLD)
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = fp.add_run("DOLCIA — 15 JUILLET 2026  •  ")
    rr.font.size = Pt(7.5)
    rr.font.color.rgb = RGBColor.from_string(MUTED)
    field(fp, "PAGE")


def page(title, kicker, intro=None):
    if len(doc.paragraphs) > 1:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(kicker.upper())
    r.font.name = "Aptos"
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(GOLD)
    h = doc.add_heading(title, 1)
    if intro:
        q = doc.add_paragraph(intro)
        q.style = styles["Normal"]
        q.paragraph_format.space_after = Pt(14)
        q.runs[0].font.size = Pt(12.5)
        q.runs[0].font.color.rgb = RGBColor.from_string(MUTED)


def bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(.55)
        p.paragraph_format.first_line_indent = Cm(-.25)
        p.add_run(item)


def quote(text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, BLACK)
    cell.margin_top = cell.margin_bottom = 220
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "Georgia"
    r.font.size = Pt(17)
    r.font.italic = True
    r.font.color.rgb = RGBColor.from_string(CREAM)


def matrix(headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, value in enumerate(headers):
        c = table.rows[0].cells[i]
        shade(c, BLACK)
        p = c.paragraphs[0]
        r = p.add_run(value)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string(GOLD)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            shade(cells[i], "F4EFE4" if ridx % 2 == 0 else "EDE5D7")
            p = cells[i].paragraphs[0]
            r = p.add_run(value)
            r.font.size = Pt(8.7)
            r.font.color.rgb = RGBColor.from_string(INK)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


# Cover
section.top_margin = Cm(.8)
section.bottom_margin = Cm(.8)
cover = doc.add_table(rows=1, cols=1)
cover.autofit = False
cover.cell(0, 0).width = Cm(17.5)
shade(cover.cell(0, 0), BLACK)
c = cover.cell(0, 0)
p = c.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("DOLCIA")
r.font.name = "Georgia"
r.font.size = Pt(39)
r.font.color.rgb = RGBColor.from_string(CREAM)
p2 = c.add_paragraph("PRODUCT BOOK MASTER v16")
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in p2.runs:
    r.font.name = "Aptos"
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(GOLD)
if HERO.exists():
    ip = c.add_paragraph()
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.add_run().add_picture(str(HERO), width=Inches(6.6))
tag = c.add_paragraph("VOS PROCHAINES ÉMOTIONS COMMENCENT ICI.")
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in tag.runs:
    r.font.name = "Georgia"
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string(CREAM)
conf = c.add_paragraph("CONFIDENTIEL — DIFFUSION CONTRÔLÉE\n15 juillet 2026 • Le Touquet, territoire pilote")
conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in conf.runs:
    r.font.name = "Aptos"
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(GOLD)

page("La promesse", "01 — Vision", "Dolcia transforme un temps libre en expérience désirable, réelle et immédiatement organisable.")
quote("Dolcia connaît le territoire, comprend mon moment et compose une expérience que je n'aurais pas imaginée seul.")
doc.add_heading("Le problème à résoudre", 2)
bullets([
    "Le touriste paie cher pour venir mais se perd dans des listes, des horaires contradictoires et des idées génériques.",
    "L'habitant connaît sa ville mais ne sait pas toujours quoi vivre aujourd'hui avec son couple, ses enfants ou ses amis.",
    "Les professionnels disposent de places, créneaux et offres invisibles au bon public au bon moment.",
])
doc.add_heading("La réponse Dolcia", 2)
bullets([
    "Un choix immense, jamais réduit artificiellement à trois cartes.",
    "Un classement qui comprend l'instant sans enfermer l'utilisateur dans un profil permanent.",
    "Un programme aux horaires réels, distinct du catalogue libre.",
    "Une exigence de vérité : une information douteuse n'entre jamais dans un agenda.",
])

page("Une architecture limpide", "02 — Produit", "Quatre espaces complémentaires. Aucun écran ne repose deux fois la même question.")
matrix(["ESPACE", "RÔLE", "PROMESSE"], [
    ("Explorer", "Voir toute la richesse disponible", "Des centaines d'idées triables, sans faux manque"),
    ("Mon programme", "Confier son temps à Dolcia", "Un agenda composé, modifiable étape par étape"),
    ("Agenda", "Garder uniquement ses choix", "Horaires éditables, réservations et rappels"),
    ("Services", "Simplifier l'expérience", "Conciergerie, garde d'enfants et services validés"),
])
doc.add_heading("Règle de navigation", 2)
bullets([
    "La date ou le séjour choisi détermine automatiquement les durées possibles.",
    "« Maintenant » commence à l'heure réelle ; « Ce soir » ne propose jamais le matin.",
    "Après-midi + soirée est un seul créneau continu ; une journée complète n'exige pas un second choix contradictoire.",
    "Explorer et Mon programme restent deux expériences visuellement et fonctionnellement distinctes.",
])

page("Le D vivant & L'Éclat", "03 — Identité", "Le D est le sceau premium. L'Éclat est la présence sensible qui aide sans infantiliser.")
if HERO.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(HERO), width=Inches(6.7))
matrix(["ÉLÉMENT", "FONCTION"], [
    ("Le D", "Emblème officiel, luxe intemporel, confiance"),
    ("L'Éclat", "Esprit abstrait dans l'étoile centrale ; deux points de lumière, pas de bouche permanente"),
    ("Écrit", "Canal précis et discret par défaut"),
    ("Voix", "Option naturelle, jamais imposée"),
])
doc.add_paragraph("L'Éclat respire doucement, s'illumine lorsqu'il comprend et devient plus chaud lorsqu'une composition est prête. Il n'est ni un robot, ni une mascotte enfantine.")

page("Le Jumeau du moment", "04 — Compréhension", "Dolcia distingue qui je suis de ce dont j'ai envie maintenant.")
matrix(["COMPRENDRE", "EXEMPLES"], [
    ("Le groupe", "Participants, âges, goûts individuels et communs"),
    ("Le moment", "Énergie, humeur, intention profonde, refus absolus"),
    ("Le réel", "Météo, trafic, distances, horaires, disponibilité"),
    ("Les contraintes", "Poussette, mobilité, animal, véhicule, rythme des enfants"),
    ("Le séjour", "Budget global restant, nuits, expériences déjà vécues"),
])
quote("Nous sommes quatre, deux adolescentes, quelque chose de vivant mais pas sportif, puis bien manger face à la mer.")
doc.add_paragraph("Une phrase peut suffire. Si elle ne suffit pas, L'Éclat pose uniquement la question qui change réellement la proposition.")

page("Le Concierge vivant", "05 — Conversation", "Une conversation courte, adaptative et utile — jamais un questionnaire déguisé.")
doc.add_heading("Exemple", 2)
matrix(["L'ÉCLAT", "OBJECTIF"], [
    ("Ce soir, retrouver de l'énergie ou décrocher complètement ?", "Comprendre le besoin dominant"),
    ("Tout vivre ensemble, ou un petit moment séparé avant de vous retrouver ?", "Comprendre la dynamique du groupe"),
    ("Une seule chose à éviter absolument ?", "Écarter l'erreur majeure"),
])
doc.add_heading("Canaux", 2)
bullets([
    "Discussion écrite disponible immédiatement.",
    "Dictée vocale utilisable comme accélérateur.",
    "Conversation voix-à-voix fluide : cible produit nécessitant une intégration serveur sécurisée.",
    "À tout moment, l'utilisateur peut ignorer le Concierge et parcourir l'intégralité du catalogue.",
])

page("La Boussole ouverte", "06 — Choix", "Les réponses hiérarchisent les idées ; elles ne ferment jamais silencieusement des portes.")
matrix(["PRINCIPE", "RÈGLE"], [
    ("Préférences", "Font remonter les idées les plus pertinentes"),
    ("Exclusions", "Uniquement explicites : activité refusée, contrainte ou impossibilité réelle"),
    ("Exploration", "Permet de sortir volontairement du profil habituel"),
    ("Transparence", "Explique pourquoi une idée remonte"),
])
doc.add_heading("La surprise maîtrisée", 2)
bullets([
    "Environ 65 % d'alignement évident, 25 % de variations intelligentes et 10 % de découvertes.",
    "Cette répartition guide le moteur ; elle ne prime jamais sur la sécurité, les horaires ou les contraintes.",
    "Une surprise n'est jamais un tirage aléatoire.",
])

page("Explorer l'immense sans se perdre", "07 — Catalogue", "Le choix doit être immense. La valeur de Dolcia est sa capacité à le rendre immédiatement désirable et lisible.")
bullets([
    "Chargement progressif par lots, recherche texte et carte vivante.",
    "Filtres naturels : accessible à pied, encore disponible, parfait avec des adolescents, moins de 90 minutes, au chaud, face à la mer.",
    "Collections éditoriales, événements du jour, grandes actualités et offres réellement exceptionnelles.",
    "Préférences visuelles choisies par l'utilisateur pour personnaliser l'ordre sans masquer le reste.",
    "Tri par pertinence du moment, distance, horaire, qualité vérifiée, prix ou nouveauté.",
])
matrix(["À 19 H", "EXEMPLES DE FAMILLES ATTENDUES"], [
    ("Dîner", "Restaurants réellement ouverts et réservables"),
    ("Sortie", "Cinéma avec séances, spectacles, bars, guinguettes, casino"),
    ("Plein air", "Balade, plage ou coucher de soleil lorsque compatible"),
    ("Événement", "Match, feu d'artifice, concert ou rendez-vous du jour confirmé"),
])

page("La composition magique", "08 — Mon programme", "Un vrai programme horaire, pas une liste renommée.")
matrix(["ÉTAPE", "CE QUE DOLCIA SÉCURISE"], [
    ("Lancement", "Heure réelle, énergie et distance"),
    ("Découverte", "Compatibilité avec le groupe et la météo"),
    ("Moment fort", "Disponibilité, réservation et fiabilité"),
    ("Respiration", "Temps de trajet et marge réaliste"),
    ("Final mémorable", "Émotion recherchée et retour possible"),
])
bullets([
    "Régénérer tout le programme ou une seule étape.",
    "Préciser librement une envie pour une étape : « autre hôtel mais avec piscine », « plus vivant », « à moins de 10 minutes ».",
    "Ajouter, déplacer, supprimer ou modifier l'horaire de chaque activité dans l'Agenda.",
    "Plan météo et plan B automatique si une activité ferme ou devient complète.",
    "Budget global du séjour, hébergement et nombre de nuits inclus — jamais un budget confondu avec une activité isolée.",
])

page("Le moteur de vérité", "09 — Fiabilité", "Dolcia préfère reconnaître un doute plutôt que fabriquer une certitude.")
matrix(["NIVEAU", "RÈGLE D'AFFICHAGE"], [
    ("Confirmé", "Horaire, adresse et disponibilité prouvés ; autorisé dans Mon programme"),
    ("Très probable", "Horaires officiels compatibles ; réservation non garantie ; signalé clairement"),
    ("À vérifier", "Visible uniquement dans Explorer, jamais imposé dans un agenda"),
    ("Refusé", "Donnée contradictoire, ancienne, hors zone ou incomplète"),
])
doc.add_heading("Contrôles universels", 2)
bullets([
    "Événement uniquement à la date et à l'heure officielles.",
    "Un lieu fermé au créneau demandé est exclu, même s'il est indiqué « ouvert actuellement » à une autre heure.",
    "Un cinéma sans séance connue n'est pas présenté comme une séance disponible.",
    "« Hôtel de ville », mairie et visite guidée ne peuvent jamais être classés comme hébergement.",
    "Un hébergement exige un type officiel compatible, une adresse et des nuits cohérentes avec le séjour.",
    "Une image doit appartenir au lieu ou être clairement marquée comme ambiance éditoriale.",
])

page("Apprendre sans enfermer", "10 — Personnalisation", "Dolcia mémorise ce qui aide et laisse l'envie du jour reprendre la main.")
bullets([
    "Favoris, idées refusées, expériences réellement vécues, sensations après la sortie et personnes présentes.",
    "Temps passé sur une fiche, proposition acceptée puis remplacée et raisons données librement.",
    "Distinction entre goûts permanents, contexte de groupe et envie ponctuelle.",
    "Contrôles pour consulter, corriger, exporter ou effacer la mémoire.",
])
matrix(["APRÈS L'EXPÉRIENCE", "SIGNAL"], [
    ("À revivre", "Souvenir fort et recommandation future"),
    ("Bien pour ce moment", "Contexte réussi sans préférence permanente"),
    ("Pas pour moi", "Déclassement expliqué"),
    ("Raconter en un mot", "Sensation libre utilisée avec prudence"),
])

page("L'intelligence familiale", "11 — Groupes", "Dolcia peut créer du commun, du séparé et des retrouvailles — uniquement lorsque tout est vérifié.")
bullets([
    "Prendre en compte les âges, rythmes, refus et goûts de chacun.",
    "Proposer une activité ensemble lorsqu'elle satisfait réellement le groupe.",
    "Suggérer des expériences parallèles parents/enfants seulement avec horaires, lieux, encadrement et trajets compatibles.",
    "Composer le point de retrouvailles : même lieu ou distance réaliste, horaire précis et réservation possible.",
    "Ne jamais assimiler automatiquement « famille » à compétition de golf ou atelier enfant.",
])
quote("Un programme réussi ne contente pas une moyenne : il orchestre les personnes présentes.")

page("Les grands rendez-vous", "12 — Temps réel", "Dolcia détecte ce qui transforme réellement une journée : sport, culture, fêtes, sorties et actualité locale.")
bullets([
    "Calendriers officiels, billetteries, organisateurs, offices de tourisme et déclarations partenaires datées.",
    "Mise en avant sur l'accueil uniquement lorsque l'événement est pertinent pour la destination et le moment.",
    "Lieux de diffusion séparés en « confirmé » et « à appeler » ; jamais un bar probable présenté comme certain.",
    "Programme cinéma enrichi des séances, avant-premières, sorties majeures et événements spéciaux lorsque la source le prouve.",
    "Les rendez-vous non sportifs bénéficient du même moteur : concerts, festivals, expositions, marchés, fêtes et phénomènes exceptionnels.",
])

page("L'écosystème partenaire invisible", "13 — Professionnels", "Le cockpit partenaire est séparé de l'application client, comme chez les meilleures plateformes de service.")
bullets([
    "Compte professionnel, justificatifs et validation humaine avant toute publication.",
    "Déclaration de disponibilité, place libérée, désistement, événement, diffusion, prix, photos et horaires.",
    "Offres flash avec prix initial, prix Dolcia inférieur, quantité restante et expiration obligatoire.",
    "Notifications push qualifiées pour remplir un créneau sans spammer le consommateur.",
    "Les packs et tarifs partenaires ne sont jamais visibles dans l'expérience client.",
    "Une entreprise payante ne passe devant qu'à pertinence consommateur égale.",
])

page("Une expérience premium 2030", "14 — Design", "Le luxe vient de la maîtrise, de la fluidité et de la confiance — pas de l'accumulation d'or.")
bullets([
    "Photographies éditoriales plein écran, fidèles au lieu et à la destination.",
    "Noir velours, champagne rare, blanc chaud et mouvement cinématographique discret.",
    "Cartes immersives, typographie mobile maîtrisée et textes toujours lisibles.",
    "Mon programme présenté comme un carnet de voyage lumineux, visuellement distinct d'Explorer.",
    "Micro-interactions utiles : compréhension, confirmation, remplacement et changement de niveau de confiance.",
    "Mode clair premium complémentaire et accessibilité sans appauvrir l'identité.",
    "Aucun bouton mystérieux, aucune information technique, aucun texte blanc sur fond blanc.",
])

page("Données & architecture", "15 — Technique", "Le Product Book décrit les responsabilités sans exposer les secrets, les clés ou la formule de classement.")
matrix(["COUCHE", "RESPONSABILITÉ"], [
    ("Sources", "Google Places, DATAtourisme, OpenAgenda, Ticketmaster, offices, partenaires"),
    ("Normalisation", "Déduplication, catégories, géolocalisation, dates, médias"),
    ("Vérité", "Contradictions, fraîcheur, confiance et exclusions"),
    ("Classement privé", "Jumeau du moment, contexte, diversité et surprise maîtrisée"),
    ("Composition", "Créneaux, trajets, budget, réservations, plans B"),
    ("Interfaces", "Client, partenaire et administration strictement séparés"),
])
bullets([
    "Clés et jetons uniquement dans les variables serveur.",
    "Le moteur central doit migrer côté serveur avant une ouverture publique.",
    "Cache, chargement progressif, observabilité et alertes sur les données contradictoires.",
    "Aucune source ne suffit seule à garantir l'exhaustivité mondiale ou la disponibilité réelle.",
])

page("Confidentialité & propriété intellectuelle", "16 — Protection", "Protéger Dolcia exige des preuves, des dépôts, de la discrétion et une architecture réellement privée.")
bullets([
    "Product Book marqué confidentiel et transmis nominativement.",
    "Accords de confidentialité avant toute remise des règles de classement, du code ou des accès partenaires.",
    "Archive e-Soleau contenant ce Product Book, le code daté, le manifeste SHA-256 et l'identité de L'Éclat.",
    "Recherche d'antériorités puis dépôts de la marque Dolcia, du D et des signes retenus.",
    "Étude d'un dépôt de dessins et modèles avant une nouvelle diffusion large des visuels.",
    "Dépôt de code privé, démonstrations protégées et secrets exclusivement serveur.",
    "Registre des destinataires, accès révoqués à la fin d'une mission et sauvegardes chiffrées.",
])
doc.add_paragraph("Une idée seule n'est pas protégée comme telle. Dolcia protège ses expressions originales, son identité, son code, ses bases structurées et son savoir-faire confidentiel. Un conseil en propriété industrielle doit valider la stratégie de dépôt.")

page("État réel & feuille de route", "17 — Exécution", "La vision est forte ; la confiance exige de distinguer ce qui fonctionne déjà de ce qui reste à industrialiser.")
matrix(["ÉTAT", "ÉLÉMENTS"], [
    ("Fonctionnel dans le MASTER", "Explorer progressif, recherche, familles, préférences, programme séparé, agenda, filtres temporels, connecteurs serveur"),
    ("Partiel / à renforcer", "Qualité et couverture des données, comptes, mémoire multi-appareils, plan B, disponibilité, programme cinéma, espace partenaire"),
    ("Cible non encore livrée", "Conversation vocale temps réel, vrai moteur IA serveur, vérité multi-source continue, réservation unifiée, couverture exhaustive"),
])
doc.add_heading("Ordre de construction", 2)
bullets([
    "1. Fiabilité universelle des dates, horaires, lieux, catégories, hébergements et images.",
    "2. Moteur serveur privé et base d'expériences normalisée Le Touquet / Côte d'Opale.",
    "3. Comptes, consentement, mémoire et apprentissage contextualisé.",
    "4. Partenaires, offres temps réel et disponibilité vérifiée.",
    "5. Concierge écrit puis voix-à-voix, expansion aux grandes villes et à la France.",
])

page("Les règles absolues", "18 — MASTER", "Ces règles priment sur la quantité, le spectaculaire et la rapidité de livraison.")
bullets([
    "Jamais de faux horaire, de fausse disponibilité, de fausse adresse ou d'image trompeuse.",
    "Jamais un événement hors date dans le programme demandé.",
    "Jamais un lieu fermé présenté comme une activité possible.",
    "Jamais une mairie classée comme hôtel, dans aucune ville du monde.",
    "Jamais un catalogue réduit silencieusement à quelques idées.",
    "Jamais une mise en avant payante au détriment de la pertinence consommateur.",
    "Jamais un doute transformé en certitude par l'IA.",
    "Toujours permettre de comprendre, ajuster, remplacer et reprendre la main.",
])
quote("Chaque instant mérite une expérience. Chaque recommandation mérite la confiance.")
doc.add_paragraph("FIN DU PRODUCT BOOK MASTER v16", style=None).alignment = WD_ALIGN_PARAGRAPH.CENTER

for sec in doc.sections:
    margins(sec)

doc.core_properties.title = "Dolcia — Product Book MASTER v16"
doc.core_properties.subject = "Vision produit, expérience, fiabilité, architecture et protection"
doc.core_properties.author = "Dolcia — document confidentiel"
doc.core_properties.keywords = "Dolcia, confidentiel, Product Book, L'Éclat, Concierge vivant"
doc.save(OUT)
print(OUT)
